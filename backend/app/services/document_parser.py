"""
Document Parser Service
Extracts, normalizes, and validates text from user-uploaded research documents:
PDF, DOCX, TXT, Markdown, and CSV.
"""

import io
import re
import hashlib
import logging
from typing import Optional
from datetime import datetime, timezone
import uuid

from app.schemas.research import ParsedDocument, Source, SourceType

logger = logging.getLogger(__name__)

# In-memory store for uploaded documents before and during research runs
_uploaded_documents: dict[str, ParsedDocument] = {}


class DocumentParser:
    """Extract and normalize text from various document formats."""

    @staticmethod
    def parse_file(filename: str, file_bytes: bytes) -> ParsedDocument:
        """
        Parse file bytes into a ParsedDocument object based on file extension.
        Handles empty, corrupt, or scanned documents without crashing.
        """
        doc_id = str(uuid.uuid4())[:8]
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        
        if not file_bytes or len(file_bytes) == 0:
            return ParsedDocument(
                document_id=doc_id,
                filename=filename,
                file_type=ext or "unknown",
                text="",
                readable=False,
                status="error",
                error="The uploaded document is empty (0 bytes).",
            )

        try:
            if ext == "pdf":
                return DocumentParser._parse_pdf(doc_id, filename, file_bytes)
            elif ext in ("docx", "doc"):
                return DocumentParser._parse_docx(doc_id, filename, file_bytes)
            elif ext in ("txt", "text"):
                return DocumentParser._parse_text(doc_id, filename, file_bytes, "txt")
            elif ext in ("md", "markdown"):
                return DocumentParser._parse_text(doc_id, filename, file_bytes, "md")
            elif ext == "csv":
                return DocumentParser._parse_csv(doc_id, filename, file_bytes)
            else:
                # Try generic text decoding
                return DocumentParser._parse_text(doc_id, filename, file_bytes, ext or "txt")
        except Exception as e:
            logger.error(f"Failed to parse document '{filename}': {e}", exc_info=True)
            return ParsedDocument(
                document_id=doc_id,
                filename=filename,
                file_type=ext or "unknown",
                text="",
                readable=False,
                status="error",
                error=f"Unable to extract text from document: {str(e)[:150]}",
            )

    @staticmethod
    def _parse_pdf(doc_id: str, filename: str, file_bytes: bytes) -> ParsedDocument:
        """Extract text from PDF using pypdf."""
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            
            if len(reader.pages) == 0:
                return ParsedDocument(
                    document_id=doc_id,
                    filename=filename,
                    file_type="pdf",
                    text="",
                    readable=False,
                    status="unreadable",
                    error="Could not extract machine-readable text from this PDF.",
                )

            extracted_pages = []
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        extracted_pages.append(f"--- Page {i + 1} ---\n" + page_text.strip())
                except Exception as pe:
                    logger.warning(f"Error extracting page {i+1} of {filename}: {pe}")

            combined_text = "\n\n".join(extracted_pages).strip()
            clean_text = DocumentParser._normalize_text(combined_text)

            # Check if meaningful machine-readable text was found
            if not clean_text or len(clean_text) < 20:
                return ParsedDocument(
                    document_id=doc_id,
                    filename=filename,
                    file_type="pdf",
                    text="",
                    readable=False,
                    status="unreadable",
                    error="Could not extract machine-readable text from this PDF.",
                )

            words = len(clean_text.split())
            return ParsedDocument(
                document_id=doc_id,
                filename=filename,
                file_type="pdf",
                char_count=len(clean_text),
                word_count=words,
                text=clean_text,
                readable=True,
                status="ready",
            )
        except Exception as e:
            logger.warning(f"pypdf extraction failed for {filename}: {e}")
            return ParsedDocument(
                document_id=doc_id,
                filename=filename,
                file_type="pdf",
                text="",
                readable=False,
                status="unreadable",
                error="Could not extract machine-readable text from this PDF.",
            )

    @staticmethod
    def _parse_docx(doc_id: str, filename: str, file_bytes: bytes) -> ParsedDocument:
        """Extract text, headings, and tables from DOCX using python-docx."""
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))

            elements = []
            # Extract paragraphs and headings
            for p in doc.paragraphs:
                p_text = p.text.strip()
                if p_text:
                    if p.style and "Heading" in p.style.name:
                        elements.append(f"\n## {p_text}\n")
                    else:
                        elements.append(p_text)

            # Extract tables
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    if any(cells):
                        table_rows.append(" | ".join(cells))
                if table_rows:
                    elements.append("\n" + "\n".join(table_rows) + "\n")

            combined_text = "\n\n".join(elements).strip()
            clean_text = DocumentParser._normalize_text(combined_text)

            if not clean_text or len(clean_text) < 10:
                return ParsedDocument(
                    document_id=doc_id,
                    filename=filename,
                    file_type="docx",
                    text="",
                    readable=False,
                    status="unreadable",
                    error="The DOCX document contains no extractable text.",
                )

            words = len(clean_text.split())
            return ParsedDocument(
                document_id=doc_id,
                filename=filename,
                file_type="docx",
                char_count=len(clean_text),
                word_count=words,
                text=clean_text,
                readable=True,
                status="ready",
            )
        except Exception as e:
            logger.warning(f"python-docx extraction failed for {filename}: {e}")
            return ParsedDocument(
                document_id=doc_id,
                filename=filename,
                file_type="docx",
                text="",
                readable=False,
                status="error",
                error=f"Failed to parse DOCX file: {str(e)[:100]}",
            )

    @staticmethod
    def _parse_text(doc_id: str, filename: str, file_bytes: bytes, file_type: str) -> ParsedDocument:
        """Extract text from plain text or markdown with encoding fallback."""
        text = ""
        for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252", "iso-8859-1"):
            try:
                text = file_bytes.decode(encoding)
                break
            except UnicodeDecodeError:
                continue

        clean_text = DocumentParser._normalize_text(text)
        if not clean_text.strip():
            return ParsedDocument(
                document_id=doc_id,
                filename=filename,
                file_type=file_type,
                text="",
                readable=False,
                status="unreadable",
                error="Document is empty or contains no readable characters.",
            )

        words = len(clean_text.split())
        return ParsedDocument(
            document_id=doc_id,
            filename=filename,
            file_type=file_type,
            char_count=len(clean_text),
            word_count=words,
            text=clean_text,
            readable=True,
            status="ready",
        )

    @staticmethod
    def _parse_csv(doc_id: str, filename: str, file_bytes: bytes) -> ParsedDocument:
        """Parse CSV safely into structured markdown table format."""
        import csv
        text = ""
        for encoding in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                text = file_bytes.decode(encoding)
                break
            except UnicodeDecodeError:
                continue

        if not text.strip():
            return ParsedDocument(
                document_id=doc_id,
                filename=filename,
                file_type="csv",
                text="",
                readable=False,
                status="unreadable",
                error="CSV file is empty.",
            )

        reader = csv.reader(io.StringIO(text))
        rows = list(reader)
        if not rows:
            return ParsedDocument(
                document_id=doc_id,
                filename=filename,
                file_type="csv",
                text="",
                readable=False,
                status="unreadable",
                error="CSV file contains no rows.",
            )

        formatted_rows = [" | ".join(row) for row in rows[:200]]  # Cap preview at 200 rows
        table_text = "\n".join(formatted_rows)
        clean_text = DocumentParser._normalize_text(table_text)

        words = len(clean_text.split())
        return ParsedDocument(
            document_id=doc_id,
            filename=filename,
            file_type="csv",
            char_count=len(clean_text),
            word_count=words,
            text=clean_text,
            readable=True,
            status="ready",
        )

    @staticmethod
    def _normalize_text(text: str) -> str:
        """Normalize extracted document text."""
        if not text:
            return ""
        # Normalize carriage returns and newlines
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        # Collapse 3+ newlines to 2
        text = re.sub(r'\n{3,}', '\n\n', text)
        # Collapse multiple horizontal whitespace to 1
        text = re.sub(r'[ \t]{2,}', ' ', text)
        # Strip null bytes and non-printable control characters
        text = "".join(ch for ch in text if ch == '\n' or ch == '\t' or ch >= ' ')
        return text.strip()

    @staticmethod
    def to_source(doc: ParsedDocument) -> Source:
        """Convert a ParsedDocument into a unified Source model for the pipeline."""
        title = doc.filename.rsplit(".", 1)[0].replace("_", " ").replace("-", " ").title()
        return Source(
            source_id=f"doc_{doc.document_id}",
            title=f"{title} ({doc.filename})",
            url=f"document://{doc.filename}",
            domain=f"Document ({doc.file_type.upper()})",
            author="User Document",
            source_type=SourceType.DOCUMENT.value,
            document_format=doc.file_type.lower(),
            content=doc.text,
            content_length=doc.char_count,
            extraction_success=doc.readable and bool(doc.text),
            extraction_error=doc.error or "",
            retrieved_at=doc.uploaded_at,
            summary=f"User-provided {doc.file_type.upper()} document: '{doc.filename}' ({doc.word_count} words).",
        )


def store_document(doc: ParsedDocument) -> None:
    """Store document in in-memory session cache."""
    _uploaded_documents[doc.document_id] = doc


def get_document(doc_id: str) -> Optional[ParsedDocument]:
    """Retrieve document by ID."""
    return _uploaded_documents.get(doc_id)


def get_all_documents() -> list[ParsedDocument]:
    """Retrieve all stored documents."""
    return list(_uploaded_documents.values())


def clear_documents() -> None:
    """Clear uploaded documents cache."""
    global _uploaded_documents
    _uploaded_documents = {}
